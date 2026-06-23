"""Processamento de anexos dos Assistentes.

Converte cada arquivo enviado em algo que o Claude entende:
- PDF  -> bloco `document` nativo (base64), o modelo lê direto.
- CSV/TXT -> texto.
- DOCX -> texto extraído (python-docx).
- XLSX -> texto tabular extraído (openpyxl), limitado em linhas.

Limites defensivos para não estourar contexto/custo nem travar o servidor.
"""
from __future__ import annotations

import base64
import io

MAX_BYTES = 15 * 1024 * 1024   # 15 MB por arquivo
MAX_FILES = 5
MAX_CHARS = 60_000             # teto de texto extraído por arquivo
MAX_ROWS = 2_000               # teto de linhas por planilha


def _texto_arquivo(nome: str, conteudo: str) -> str:
    return f"### Arquivo: {nome}\n{conteudo[:MAX_CHARS]}\n"


def processar_anexo(filename: str, data: bytes) -> tuple[list[dict], str]:
    """Devolve (blocos_document, texto_extraido) para um arquivo.

    blocos_document: blocos de conteúdo p/ a API (PDF). texto_extraido: texto a
    anexar à mensagem (CSV/DOCX/XLSX). Tipos não suportados são ignorados.
    """
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    if ext == "pdf":
        bloco = {
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf",
                       "data": base64.b64encode(data).decode("ascii")},
        }
        return [bloco], ""

    if ext in ("csv", "txt"):
        return [], _texto_arquivo(filename, data.decode("utf-8", errors="replace"))

    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            partes = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            # também extrai tabelas simples
            for t in doc.tables:
                for row in t.rows:
                    partes.append(" | ".join(c.text for c in row.cells))
            return [], _texto_arquivo(filename, "\n".join(partes))
        except Exception:
            return [], f"### Arquivo: {filename}\n(não foi possível ler o Word)\n"

    if ext in ("xlsx", "xlsm"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            linhas: list[str] = []
            for ws in wb.worksheets:
                linhas.append(f"-- Planilha: {ws.title} --")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= MAX_ROWS:
                        linhas.append("… (linhas restantes truncadas)")
                        break
                    linhas.append(", ".join("" if c is None else str(c) for c in row))
            return [], _texto_arquivo(filename, "\n".join(linhas))
        except Exception:
            return [], f"### Arquivo: {filename}\n(não foi possível ler a planilha)\n"

    # Tipo não suportado — ignora silenciosamente (a UI já filtra os aceitos).
    return [], ""


def extrair_texto(filename: str, data: bytes) -> str:
    """Extrai SÓ texto de um arquivo (inclui PDF via pypdf) — para o contexto
    persistido de projetos (injetado a todos os agentes como texto)."""
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    if ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            txt = "\n".join((p.extract_text() or "") for p in reader.pages)
            return txt[:MAX_CHARS] or "(PDF sem texto extraível — pode ser digitalizado/imagem)"
        except Exception:
            return "(não foi possível ler o PDF)"

    # CSV/TXT/DOCX/XLSX reaproveitam o extrator de texto já existente.
    _, txt = processar_anexo(filename, data)
    return txt
